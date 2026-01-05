import streamlit as st
import numpy as np
import random
import time
import base64
import io
import json
import hashlib
st.title("SUDUKU GENERATOR")
# Try to import docx, but provide fallback if not installed
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Set page configuration
st.set_page_config(
    page_title="Sudoku Game Generator",
    page_icon="🔢",
    layout="wide"
)

# Custom CSS with all fixes
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #1E3A8A;
        text-align: center;
        margin-bottom: 1rem;
    }
    .sudoku-container {
        background-color: white;
        padding: 10px;
        border-radius: 10px;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        margin: 0 auto;
        max-width: 500px;
    }
   # .timer {
    #    font-size: 1.5rem;
      #  font-weight: bold;
      #  color: #1E3A8A;
      #  text-align: center;
       # padding: 10px;
       # background-color: #F0F9FF;
       # border-radius: 8px;
       # margin-bottom: 10px;
   # }
    .number-buttons {
        display: flex;
        justify-content: center;
        gap: 8px;
        margin: 15px 0;
        flex-wrap: wrap;
    }
    .export-buttons {
        display: flex;
        justify-content: center;
        gap: 10px;
        margin: 15px 0;
        flex-wrap: wrap;
    }
    
    /* SUDOKU GRID STYLING */
    .sudoku-cell {
        width: 45px !important;
        height: 45px !important;
        min-width: 45px !important;
        min-height: 45px !important;
        padding: 0 !important;
        margin: 0 !important;
        font-size: 26px !important;
        font-weight: bold !important;
        border: 1px solid #94A3B8 !important;
        border-radius: 0 !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        line-height: 1 !important;
    }
    
    [data-testid="column"] {
        padding: 0 !important;
        margin: 0 !important;
        gap: 0 !important;
    }
    
    [data-testid="stHorizontalBlock"] {
        gap: 0 !important;
        margin: 0 !important;
        padding: 0 !important;
    }
    
    .thick-right {
        border-right: 3px solid #1E293B !important;
    }
    
    .thick-bottom {
        border-bottom: 3px solid #1E293B !important;
    }
    
    .given-cell {
        background-color: #F1F5F9 !important;
        color: #1E293B !important;
        cursor: default !important;
    }
    
    .user-cell {
        background-color: white !important;
        color: #1E40AF !important;
    }
    
    @keyframes blink {
        0% { background-color: #BFDBFE; }
        50% { background-color: #93C5FD; }
        100% { background-color: #BFDBFE; }
    }
    
    .selected-cell {
        animation: blink 1s infinite !important;
        border: 2px solid #3B82F6 !important;
        color: #1E40AF !important;
    }
    
    .error-cell {
        background-color: #FEE2E2 !important;
        color: #DC2626 !important;
    }
    
    .hint-cell {
        background-color: #FEF3C7 !important;
        color: #D97706 !important;
    }
    
    .num-btn > button {
        width: 45px !important;
        height: 45px !important;
        font-size: 24px !important;
        font-weight: bold !important;
    }
    
    .clear-btn > button {
        background-color: #FEE2E2 !important;
        color: #DC2626 !important;
        border-color: #FCA5A5 !important;
        width: 45px !important;
        height: 45px !important;
        font-size: 20px !important;
    }
    
    .instructions {
        font-size: 14px;
        color: #64748B;
        text-align: center;
        margin-top: 10px;
    }
</style>
""", unsafe_allow_html=True)

class SudokuGame:
    def __init__(self, difficulty='medium'):
        self.difficulty = difficulty
        self.board = None
        self.solution = None
        self.user_board = None
        self.start_time = None
        self.end_time = None
        self.hints_used = 0
        self.errors = 0
        self.puzzle_id = None
        self.generate_new_puzzle()
    
    def generate_new_puzzle(self):
        # Generate a complete Sudoku solution
        self.solution = self.generate_solution()
        
        # Create puzzle by removing numbers based on difficulty
        self.board = np.copy(self.solution)
        self.user_board = np.copy(self.board)
        
        # Number of cells to remove based on difficulty
        difficulty_levels = {
            'easy': 35,
            'medium': 45,
            'hard': 50,
            'expert': 55
        }
        
        cells_to_remove = difficulty_levels.get(self.difficulty, 45)
        
        # Remove numbers randomly
        all_positions = [(i, j) for i in range(9) for j in range(9)]
        random.shuffle(all_positions)
        
        removed = 0
        for i, j in all_positions:
            if removed >= cells_to_remove:
                break
            
            temp = self.board[i][j]
            self.board[i][j] = 0
            self.user_board[i][j] = 0
            
            if self.has_unique_solution():
                removed += 1
            else:
                self.board[i][j] = temp
                self.user_board[i][j] = temp
        
        self.start_time = time.time()
        self.end_time = None
        self.hints_used = 0
        self.errors = 0
        self.generate_puzzle_id()
    
    def generate_puzzle_id(self):
        """Generate unique ID for sharing puzzles"""
        puzzle_data = {
            'board': self.board.tolist(),
            'difficulty': self.difficulty,
            'timestamp': time.time()
        }
        puzzle_str = json.dumps(puzzle_data, sort_keys=True)
        self.puzzle_id = hashlib.md5(puzzle_str.encode()).hexdigest()[:8]
    
    def generate_solution(self):
        board = np.zeros((9, 9), dtype=int)
        
        def fill_board(board, row=0, col=0):
            if row == 9:
                return True
            if col == 9:
                return fill_board(board, row + 1, 0)
            if board[row][col] != 0:
                return fill_board(board, row, col + 1)
            
            numbers = list(range(1, 10))
            random.shuffle(numbers)
            
            for num in numbers:
                if self.is_valid(board, row, col, num):
                    board[row][col] = num
                    if fill_board(board, row, col + 1):
                        return True
                    board[row][col] = 0
            return False
        
        # Fill diagonal boxes first
        for i in range(0, 9, 3):
            numbers = list(range(1, 10))
            random.shuffle(numbers)
            for r in range(3):
                for c in range(3):
                    board[i + r][i + c] = numbers.pop()
        
        fill_board(board)
        return board
    
    def is_valid(self, board, row, col, num):
        if num in board[row]:
            return False
        if num in board[:, col]:
            return False
        box_row = (row // 3) * 3
        box_col = (col // 3) * 3
        if num in board[box_row:box_row+3, box_col:box_col+3]:
            return False
        return True
    
    def has_unique_solution(self):
        board_copy = np.copy(self.board)
        
        def count_solutions(board, count=0):
            for i in range(9):
                for j in range(9):
                    if board[i][j] == 0:
                        for num in range(1, 10):
                            if self.is_valid(board, i, j, num):
                                board[i][j] = num
                                count = count_solutions(board, count)
                                if count > 1:
                                    return count
                                board[i][j] = 0
                        return count
            return count + 1
        
        return count_solutions(board_copy) == 1
    
    def get_hint(self):
        empty_cells = []
        for i in range(9):
            for j in range(9):
                if self.user_board[i][j] == 0:
                    empty_cells.append((i, j))
        
        if empty_cells:
            row, col = random.choice(empty_cells)
            self.user_board[row][col] = self.solution[row][col]
            self.hints_used += 1
            return row, col
        return None
    
    def check_solution(self):
        for i in range(9):
            for j in range(9):
                if self.user_board[i][j] != self.solution[i][j]:
                    return False
        self.end_time = time.time()
        return True
    
    def check_errors(self):
        errors = []
        for i in range(9):
            for j in range(9):
                if self.user_board[i][j] != 0:
                    temp = self.user_board[i][j]
                    self.user_board[i][j] = 0
                    
                    if not self.is_valid(self.user_board, i, j, temp):
                        errors.append((i, j))
                    
                    self.user_board[i][j] = temp
        return errors
    
    def get_elapsed_time(self):
        if self.start_time is None:
            return 0
        if self.end_time is not None:
            return int(self.end_time - self.start_time)
        return int(time.time() - self.start_time)
    
    def format_time(self, seconds):
        minutes = seconds // 60
        seconds = seconds % 60
        return f"{minutes:02d}:{seconds:02d}"
    
    # ADDED: is_complete method
    def is_complete(self):
        """Check if all cells are filled (not necessarily correctly)"""
        return 0 not in self.user_board
    
    def export_to_word(self, include_solution=False):
        """Export puzzle to Word document"""
        if not DOCX_AVAILABLE:
            # Return a simple text file if docx is not available
            return self.export_to_text_file(include_solution)
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Sudoku Puzzle', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Difficulty and info
        doc.add_paragraph(f'Difficulty: {self.difficulty.title()}')
        doc.add_paragraph(f'Puzzle ID: {self.puzzle_id}')
        doc.add_paragraph(f'Date: {time.strftime("%Y-%m-%d %H:%M")}')
        doc.add_paragraph()
        
        # Create puzzle grid
        doc.add_heading('Puzzle Grid', level=1)
        
        # Create table for puzzle
        table = doc.add_table(rows=9, cols=9)
        table.style = 'Table Grid'
        
        # Fill table with puzzle numbers
        for i in range(9):
            row_cells = table.rows[i].cells
            for j in range(9):
                cell = row_cells[j]
                cell.text = str(self.board[i][j]) if self.board[i][j] != 0 else ""
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Style for given numbers (bold)
                if self.board[i][j] != 0:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(14)
                else:
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
        
        doc.add_paragraph()
        
        # Add solution if requested
        if include_solution:
            doc.add_heading('Solution', level=1)
            
            # Create table for solution
            sol_table = doc.add_table(rows=9, cols=9)
            sol_table.style = 'Table Grid'
            
            for i in range(9):
                row_cells = sol_table.rows[i].cells
                for j in range(9):
                    cell = row_cells[j]
                    cell.text = str(self.solution[i][j])
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(14)
        
        # Add instructions
        doc.add_page_break()
        doc.add_heading('How to Solve Sudoku', level=1)
        doc.add_paragraph('''Sudoku Rules:
1. Each row must contain the numbers 1-9 exactly once
2. Each column must contain the numbers 1-9 exactly once
3. Each 3x3 box must contain the numbers 1-9 exactly once

Tips:
• Start with rows/columns that have the most numbers
• Look for numbers that can only go in one cell
• Use pencil marks for possible numbers
• Take your time and be patient!''')
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        return doc_bytes
    
    def export_to_text_file(self, include_solution=False):
        """Export puzzle as text file with better formatting"""
        text = f"SUDOKU PUZZLE\n"
        text += "=" * 40 + "\n"
        text += f"Difficulty: {self.difficulty.title()}\n"
        text += f"Puzzle ID: {self.puzzle_id}\n"
        text += f"Date: {time.strftime('%Y-%m-%d %H:%M')}\n"
        text += "=" * 40 + "\n\n"
        
        # Create ASCII art grid
        text += "PUZZLE GRID:\n"
        text += "╔═══════╦═══════╦═══════╗\n"
        
        for i in range(9):
            if i % 3 == 0 and i > 0:
                text += "╠═══════╬═══════╬═══════╣\n"
            
            text += "║ "
            for j in range(9):
                if j % 3 == 0 and j > 0:
                    text += "║ "
                
                value = self.board[i][j]
                text += str(value) if value != 0 else "·"
                text += " "
            
            text += "║\n"
        
        text += "╚═══════╩═══════╩═══════╝\n\n"
        
        if include_solution:
            text += "SOLUTION:\n"
            text += "╔═══════╦═══════╦═══════╗\n"
            
            for i in range(9):
                if i % 3 == 0 and i > 0:
                    text += "╠═══════╬═══════╬═══════╣\n"
                
                text += "║ "
                for j in range(9):
                    if j % 3 == 0 and j > 0:
                        text += "║ "
                    
                    text += str(self.solution[i][j]) + " "
                
                text += "║\n"
            
            text += "╚═══════╩═══════╩═══════╝\n\n"
        
        # Add instructions
        text += "INSTRUCTIONS:\n"
        text += "-" * 40 + "\n"
        text += "1. Fill each row with numbers 1-9 (no repeats)\n"
        text += "2. Fill each column with numbers 1-9 (no repeats)\n"
        text += "3. Fill each 3x3 box with numbers 1-9 (no repeats)\n"
        text += "4. Use pencil marks for possible numbers\n"
        text += "5. Start with rows/columns with most given numbers\n\n"
        
        text += "Share this puzzle with friends using Puzzle ID: " + self.puzzle_id
        
        return text
    
    def export_to_simple_text(self):
        """Simple text export without fancy formatting"""
        text = f"Sudoku Puzzle - {self.difficulty.title()}\n"
        text += f"Puzzle ID: {self.puzzle_id}\n"
        text += f"Date: {time.strftime('%Y-%m-%d %H:%M')}\n"
        text += "=" * 40 + "\n\n"
        
        # Create simple grid
        for i in range(9):
            if i % 3 == 0 and i > 0:
                text += "------+-------+------\n"
            
            for j in range(9):
                if j % 3 == 0 and j > 0:
                    text += "| "
                
                value = self.board[i][j]
                text += str(value) if value != 0 else "."
                text += " "
            
            text += "\n"
        
        text += "\nInstructions:\n"
        text += "1. Fill each row with numbers 1-9 (no repeats)\n"
        text += "2. Fill each column with numbers 1-9 (no repeats)\n"
        text += "3. Fill each 3x3 box with numbers 1-9 (no repeats)\n"
        
        return text
    
    def get_shareable_data(self):
        """Get data for sharing"""
        return {
            'puzzle_id': self.puzzle_id,
            'difficulty': self.difficulty,
            'board': self.board.tolist(),
            'solution': self.solution.tolist(),
            'timestamp': time.time(),
            'share_text': f"""🎮 Sudoku Challenge!

Difficulty: {self.difficulty.title()}
Puzzle ID: {self.puzzle_id}
Date: {time.strftime('%Y-%m-%d')}

Can you solve it? Share your time!"""
        }

# Initialize session state
if 'game' not in st.session_state:
    st.session_state.game = SudokuGame()
if 'show_errors' not in st.session_state:
    st.session_state.show_errors = False
if 'show_hint' not in st.session_state:
    st.session_state.show_hint = False
if 'hint_cell' not in st.session_state:
    st.session_state.hint_cell = None
if 'game_over' not in st.session_state:
    st.session_state.game_over = False
if 'selected_cell' not in st.session_state:
    st.session_state.selected_cell = None
if 'last_update' not in st.session_state:
    st.session_state.last_update = time.time()
if 'show_export_options' not in st.session_state:
    st.session_state.show_export_options = False
if 'last_blink' not in st.session_state:
    st.session_state.last_blink = time.time()

# Main app
st.markdown('<div class="main-header">🔢 SUDUKU BELOW</div>', unsafe_allow_html=True)

# Sidebar
with st.sidebar:
    st.header("🎮 Game Controls")
    
    # Difficulty - WITH UNIQUE KEY ADDED
    difficulty = st.selectbox(
        "Difficulty",
        ["easy", "medium", "hard", "expert"],
        index=1,
        key="difficulty_selector"  # Added unique key
    )
    
    # New game
    if st.button("🔄 New Game", type="primary", use_container_width=True):
        st.session_state.game = SudokuGame(difficulty)
        st.session_state.show_errors = False
        st.session_state.show_hint = False
        st.session_state.hint_cell = None
        st.session_state.game_over = False
        st.session_state.selected_cell = None
        st.success(f"New {difficulty} game started!")
        st.rerun()
    
    st.divider()
    
    # Export/Share section
    st.subheader("📤 Export & Share")
    
    if st.button("💾 Export Puzzle", use_container_width=True):
        st.session_state.show_export_options = True
        st.rerun()
    
    if st.button("📱 Share Puzzle", use_container_width=True):
        share_data = st.session_state.game.get_shareable_data()
        
        # Create shareable text area
        st.text_area("Copy this to share:", share_data['share_text'], height=150)
        
        # Show puzzle ID
        st.info("📱 Puzzle ID (copy this):")
        st.code(share_data['puzzle_id'])
        
        # Show simple grid for preview
        st.write("Puzzle preview:")
        board = st.session_state.game.board
        preview_text = ""
        for i in range(9):
            for j in range(9):
                value = board[i][j]
                preview_text += str(value) if value != 0 else "."
                preview_text += " "
                if j % 3 == 2 and j < 8:
                    preview_text += "| "
            preview_text += "\n"
            if i % 3 == 2 and i < 8:
                preview_text += "------+-------+------\n"
        st.code(preview_text)
    
    st.divider()
    
    # Game helpers
    st.subheader("Helpers")
    
    if st.button("💡 Get Hint", use_container_width=True):
        if not st.session_state.game_over:
            hint_cell = st.session_state.game.get_hint()
            if hint_cell:
                st.session_state.show_hint = True
                st.session_state.hint_cell = hint_cell
                st.session_state.selected_cell = hint_cell
                st.success(f"Hint given! ({st.session_state.game.hints_used} used)")
            else:
                st.warning("No empty cells!")
            st.rerun()
    
    if st.button("🔍 Check Errors", use_container_width=True):
        if not st.session_state.game_over:
            st.session_state.show_errors = True
            errors = st.session_state.game.check_errors()
            if errors:
                st.warning(f"Found {len(errors)} errors!")
            else:
                st.success("No errors!")
            st.rerun()
    
    if st.button("✅ Check Solution", type="primary", use_container_width=True):
        if not st.session_state.game_over:
            if st.session_state.game.check_solution():
                st.session_state.game_over = True
                st.balloons()
                st.success("🎉 Puzzle solved!")
            else:
                st.error("Solution incorrect!")
            st.rerun()
    
    st.divider()
    
    # Stats
    if st.session_state.game:
        st.subheader("📊 Stats")
        game = st.session_state.game
        st.write(f"**Time:** {game.format_time(game.get_elapsed_time())}")
        st.write(f"**Hints:** {game.hints_used}")
        filled = 81 - np.count_nonzero(game.user_board == 0)
        st.write(f"**Progress:** {filled}/81")
        st.write(f"**Difficulty:** {game.difficulty.title()}")
        st.write(f"**Puzzle ID:** `{game.puzzle_id}`")

# Main game area
col1, col2, col3 = st.columns([1, 2, 1])

with col2:
    # Timer
    if not st.session_state.game_over:
        elapsed = st.session_state.game.get_elapsed_time()
        st.markdown(f'<div class="timer">⏱️ {st.session_state.game.format_time(elapsed)}</div>', unsafe_allow_html=True)
    
    # Export options (if shown)
    if st.session_state.show_export_options:
        st.info("📤 Export Options")
        col_export1, col_export2, col_export3 = st.columns(3)
        
        with col_export1:
            # Export as Formatted Text
            text_content = st.session_state.game.export_to_text_file(include_solution=False)
            st.download_button(
                label="📝 Text (Formatted)",
                data=text_content,
                file_name=f"sudoku_{st.session_state.game.puzzle_id}.txt",
                mime="text/plain",
                use_container_width=True
            )
        
        with col_export2:
            # Export as Text with Solution
            text_content_sol = st.session_state.game.export_to_text_file(include_solution=True)
            st.download_button(
                label="📝 Text (with Solution)",
                data=text_content_sol,
                file_name=f"sudoku_solution_{st.session_state.game.puzzle_id}.txt",
                mime="text/plain",
                use_container_width=True
            )
        
        with col_export3:
            # Export as Simple Text
            simple_text = st.session_state.game.export_to_simple_text()
            st.download_button(
                label="📄 Simple Text",
                data=simple_text,
                file_name=f"sudoku_simple_{st.session_state.game.puzzle_id}.txt",
                mime="text/plain",
                use_container_width=True
            )
        
        # Word export if available
        if DOCX_AVAILABLE:
            col_export4, col_export5 = st.columns(2)
            
            with col_export4:
                # Export as Word (Puzzle only)
                word_bytes = st.session_state.game.export_to_word(include_solution=False)
                st.download_button(
                    label="📄 Word (Puzzle)",
                    data=word_bytes,
                    file_name=f"sudoku_{st.session_state.game.puzzle_id}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            
            with col_export5:
                # Export as Word (with solution)
                word_bytes_sol = st.session_state.game.export_to_word(include_solution=True)
                st.download_button(
                    label="📄 Word (Puzzle+Solution)",
                    data=word_bytes_sol,
                    file_name=f"sudoku_solution_{st.session_state.game.puzzle_id}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
        else:
            st.info("💡 Install python-docx for Word export: `pip install python-docx`")
        
        if st.button("Close Export Options"):
            st.session_state.show_export_options = False
            st.rerun()
        
        st.divider()
    
    # Sudoku grid
    game = st.session_state.game
    board = game.user_board
    original = game.board
    
    st.markdown('<div class="sudoku-container">', unsafe_allow_html=True)
    
    # Create grid with st.columns for better control
    for i in range(9):
        cols = st.columns(9)
        for j in range(9):
            with cols[j]:
                # Determine cell value
                cell_value = original[i][j] if original[i][j] != 0 else board[i][j]
                display_text = str(cell_value) if cell_value != 0 else ""
                
                # Determine CSS classes
                css_classes = ["sudoku-cell"]
                
                # Thick borders for 3x3 boxes
                if j % 3 == 2 and j < 8:
                    css_classes.append("thick-right")
                if i % 3 == 2 and i < 8:
                    css_classes.append("thick-bottom")
                
                # Check cell type
                if original[i][j] != 0:
                    # Given cell
                    css_classes.append("given-cell")
                else:
                    # User cell
                    css_classes.append("user-cell")
                    
                    # Check if selected (with blinking)
                    if st.session_state.selected_cell == (i, j):
                        css_classes.append("selected-cell")
                    
                    # Check for errors
                    if st.session_state.show_errors:
                        errors = game.check_errors()
                        if (i, j) in errors:
                            css_classes.append("error-cell")
                    
                    # Check for hint
                    if st.session_state.show_hint and st.session_state.hint_cell == (i, j):
                        css_classes.append("hint-cell")
                
                # Create a clickable button for the cell
                cell_key = f"cell_{i}_{j}_{game.puzzle_id}"
                
                # Use st.button directly with custom class
                if st.button(
                    display_text, 
                    key=cell_key,
                    use_container_width=True
                ):
                    if original[i][j] == 0:
                        st.session_state.selected_cell = (i, j)
                        st.session_state.last_update = time.time()
                        st.session_state.last_blink = time.time()
                        st.rerun()
                
                # Apply CSS classes via JavaScript
                css_class_str = " ".join(css_classes)
                st.markdown(f"""
                <script>
                    // Apply CSS classes to the button
                    const button = document.querySelector('button[data-testid="stButton"][data-test-state="{cell_key}"]');
                    if (button) {{
                        button.className = '{css_class_str}';
                        const parentDiv = button.parentElement;
                        if (parentDiv) {{
                            parentDiv.className = '{css_class_str}';
                        }}
                    }}
                </script>
                """, unsafe_allow_html=True)
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Auto-refresh for blinking animation
    current_time = time.time()
    if (current_time - st.session_state.last_blink > 0.5 and 
        st.session_state.selected_cell and not st.session_state.game_over):
        st.session_state.last_blink = current_time
        st.rerun()
    
    # Number buttons
    st.markdown('<div class="number-buttons">', unsafe_allow_html=True)
    
    # Create number buttons 1-9
    num_cols = st.columns(10)  # 9 numbers + 1 clear button
    
    for num in range(1, 10):
        with num_cols[num-1]:
            if st.button(str(num), key=f"num_{num}_{game.puzzle_id}", use_container_width=True):
                if st.session_state.selected_cell:
                    i, j = st.session_state.selected_cell
                    if original[i][j] == 0:
                        game.user_board[i][j] = num
                        st.session_state.last_update = time.time()
                        # Move to next cell automatically
                        found_next = False
                        for next_j in range(j+1, 9):
                            if original[i][next_j] == 0:
                                st.session_state.selected_cell = (i, next_j)
                                found_next = True
                                break
                        if not found_next:
                            for next_i in range(i+1, 9):
                                for next_j in range(9):
                                    if original[next_i][next_j] == 0:
                                        st.session_state.selected_cell = (next_i, next_j)
                                        found_next = True
                                        break
                                if found_next:
                                    break
                        st.rerun()
    
    # Clear button
    with num_cols[9]:
        if st.button("✕", key=f"clear_{game.puzzle_id}", type="secondary", use_container_width=True):
            if st.session_state.selected_cell:
                i, j = st.session_state.selected_cell
                if original[i][j] == 0:
                    game.user_board[i][j] = 0
                    st.session_state.last_update = time.time()
                    st.rerun()
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Quick Export buttons
    st.markdown('<div class="export-buttons">', unsafe_allow_html=True)
    
    col_exp1, col_exp2 = st.columns(2)
    
    with col_exp1:
        if st.button("📥 Download Puzzle", key="download_puzzle", use_container_width=True):
            st.session_state.show_export_options = True
            st.rerun()
    
    with col_exp2:
        if st.button("📤 Share Puzzle ID", key="share_puzzle_id", use_container_width=True):
            share_data = game.get_shareable_data()
            st.info(f"Puzzle ID: **{share_data['puzzle_id']}**")
            st.code(f"Sudoku Puzzle ID: {share_data['puzzle_id']}")
            st.caption("Share this ID with friends to play the same puzzle!")
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Instructions
   # if st.session_state.selected_cell:
 #       i, j = st.session_state.selected_cell
   #     st.info(f"🔵 **Selected: Row {i+1}, Column {j+1}** (should blink)")
     #   if original[i][j] == 0:
 #           st.caption("Click a number above to enter it. Cell will blink until you select another.")
#    else:
#        st.info("👆 Click on any empty cell to select it (it will blink), then click a number")
    
    # Game status - NOW USING THE is_complete() METHOD
    if st.session_state.game_over:
        st.success(f"🎉 **Solved!** Time: {game.format_time(game.get_elapsed_time())}")
    elif game.is_complete():
        st.warning("✅ All cells filled! Click 'Check Solution' to verify.")

# Bottom section
st.divider()

col4, col5, col6 = st.columns(3)

with col4:
    # Progress
    filled = 81 - np.count_nonzero(game.user_board == 0)
    progress = filled / 81
    st.subheader("Progress")
    st.progress(progress)
    st.caption(f"{filled}/81 cells ({progress*100:.0f}%)")

with col5:
    # Quick actions
    st.subheader("Quick Actions")
    
    if st.button("🗑️ Clear All", key="clear_all", use_container_width=True):
        for i in range(9):
            for j in range(9):
                if original[i][j] == 0:
                    game.user_board[i][j] = 0
        st.session_state.show_errors = False
        st.session_state.selected_cell = None
        st.success("Cleared all your numbers!")
        st.rerun()
    
    #if st.button("📋 Copy Puzzle ID", key="copy_puzzle_id", use_container_width=True):
      #  st.code(game.puzzle_id)
       #   st.success("Puzzle ID copied!")

with col6:
    # Puzzle info
    st.subheader("Puzzle Info")
    #st.write(f"**ID:** `{game.puzzle_id}`")
    st.write(f"**Difficulty:** {game.difficulty.title()}")
    given_count = np.count_nonzero(game.board != 0)
    st.write(f"**Given numbers:** {given_count}")
    st.write(f"**Empty cells:** {81 - given_count}")
st.write("developed by Subramanian Ramajayam")


# Footer
st.divider()
st.markdown(f'<div class="instructions">🔢 Sudoku Game • Puzzle ID: {game.puzzle_id} • Export to multiple formats • Share with friends!</div>', unsafe_allow_html=True)
